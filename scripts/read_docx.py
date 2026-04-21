from docx import Document

for fname in [
    "/proj/Documentation_Technique_Pipeline_KOMOSSI_Sosso_v6.docx",
    "/proj/Memoire_Master2_IBDIA_KOMOSSI_Sosso_v4.docx"
]:
    print(f"\n{'='*80}")
    print(f"FILE: {fname}")
    print('='*80)
    doc = Document(fname)
    for para in doc.paragraphs:
        if para.text.strip():
            print(f"[{para.style.name}] {para.text}")
    for i, table in enumerate(doc.tables):
        print(f"[TABLE {i}]")
        for row in table.rows:
            cells = [c.text.strip().replace('\n',' ')[:60] for c in row.cells]
            print("  | ".join(cells))
